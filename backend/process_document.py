import sys
import torch
import re
import os
from docx import Document
from docx.shared import RGBColor
from transformers import T5ForConditionalGeneration, T5Tokenizer

# Load trained GEC model
MODEL_PATH = "C:/Users/bibis/PycharmProjects/GEC/T5_model_05"

print("[INFO] Loading tokenizer and model...")
tokenizer = T5Tokenizer.from_pretrained(MODEL_PATH)
model = T5ForConditionalGeneration.from_pretrained(MODEL_PATH)
print("[INFO] Model loaded successfully!")

def split_into_sentences(text):
    sentence_endings = r'(?<=[.!?]) +'
    sentences = re.split(sentence_endings, text)
    return [s.strip() for s in sentences if s.strip()]

def correct_grammar(input_text, num_return_sequences=1):
    batch = tokenizer([input_text], truncation=True, padding='max_length', max_length=128, return_tensors="pt").to("cpu")
    translated = model.generate(**batch, max_length=256, min_length=15, do_sample=True, num_beams=5,  num_return_sequences=num_return_sequences, temperature=0.9)
    tgt_text = tokenizer.batch_decode(translated, skip_special_tokens=True)
    return tgt_text


def process_docx(input_docx, output_docx):
    """Extracts text from Word, corrects grammar, and creates a new corrected Word file."""
    print(f"[INFO] Opening Word document: {input_docx}")
    doc = Document(input_docx)
    
    corrected_doc = Document()
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            corrected_doc.add_paragraph("") 
            continue

        sentences_original = split_into_sentences(text)
        sentences_corrected = [correct_grammar(sentence)[0] for sentence in sentences_original]

        corrected_para = corrected_doc.add_paragraph()

        for original, corrected in zip(sentences_original, sentences_corrected):
            words_original = original.split()
            words_corrected = corrected.split()
            
            i, j = 0, 0
            while i < len(words_original) and j < len(words_corrected):
                if words_original[i] == words_corrected[j]:  
                    corrected_para.add_run(words_corrected[j] + " ")  
                else:
                    run = corrected_para.add_run(words_corrected[j] + " ")
                    run.font.color.rgb = RGBColor(255, 0, 0)  
                i += 1
                j += 1

            while j < len(words_corrected): 
                run = corrected_para.add_run(words_corrected[j] + " ")
                run.font.color.rgb = RGBColor(255, 0, 0) 
                j += 1

    if os.path.exists(output_docx):
        os.remove(output_docx)

    corrected_doc.save(output_docx)
    print(f"[INFO] Corrected document saved as: {output_docx}")



if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("[ERROR] No input file provided. Usage: python process_docx.py <input.docx>")
        sys.exit(1)

    input_file = sys.argv[1]
    

    input_filename = os.path.splitext(os.path.basename(input_file))[0]
    
    output_file = f"{input_filename}_corrected.docx"
    output_file = output_file.encode("utf-8").decode("utf-8")
    
    try:
        process_docx(input_file, output_file)
        print("[SUCCESS] Word document processing completed successfully!")
    except Exception as e:
        print(f"[ERROR] An error occurred: {e}") 
