#import sys
#import re
#import json
#import spacy
#from collections import Counter
#from docx import Document

#def extract_text_from_docx(file_path):
#    doc = Document(file_path)
#    return "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])

#def split_into_sentences(text):
#    return re.split(r'(?<=[.!?]) +', text)

#def analyze_text(text):
#    nlp = spacy.load("ro_core_news_sm")
#    doc = nlp(text)

#    sentences = split_into_sentences(text)
#    words = re.findall(r'\b\w+\b', text.lower())

#    word_count = len(words)
#    sentence_count = len(sentences)
#    avg_sentence_length = word_count / sentence_count if sentence_count else 0

#    # Sentence length distribution
#    sentence_lengths = [len(re.findall(r'\b\w+\b', sentence)) for sentence in sentences]
#    sentence_length_distribution = Counter(sentence_lengths).most_common()

#    # Dicendi verbs (custom list)
#    dicendi_verbs = {"conversa", "dialoga", "vorbi", "discuta", "bârfi", "dezbate", "delibera"}
#    dicendi_count = sum(1 for token in doc if token.lemma_ in dicendi_verbs and token.pos_ == "VERB")

#    # Abstract vs concrete words (mock logic)
#    abstract_terms = {"iubire", "frică", "libertate", "gândire", "idee", "emoție"}
#    concrete_count = 0
#    abstract_count = 0
#    for token in doc:
#        if token.is_alpha and token.pos_ in {"NOUN", "ADJ", "VERB"}:
#            if token.lemma_.lower() in abstract_terms:
#                abstract_count += 1
#            else:
#                concrete_count += 1

#    # Sentence types
#    type_counts = {"declarative": 0, "interogative": 0, "exclamative": 0}
#    for sentence in sentences:
#        if sentence.endswith("?"):
#            type_counts["interogative"] += 1
#        elif sentence.endswith("!"):
#            type_counts["exclamative"] += 1
#        else:
#            type_counts["declarative"] += 1

#    # Most common words (excluding stopwords and punctuation)
#    filtered_words = [
#        token.text.lower() for token in doc
#        if token.pos_ not in {"DET", "ADP", "CCONJ", "SCONJ", "PART", "INTJ", "PUNCT", "AUX"}
#    ]
#    most_common_words = Counter(filtered_words).most_common(5)

#    # Vocabulary analysis
#    unique_words = set(words)
#    uniqueness_ratio = round(len(unique_words) / word_count, 2) if word_count else 0

#    filler_count = sum(1 for token in doc if token.pos_ in {"ADP", "AUX", "SCONJ", "CCONJ"})
#    content_count = word_count - filler_count

#    return {
#        "word_count": word_count,
#        "sentence_count": sentence_count,
#        "avg_sentence_length": round(avg_sentence_length, 2),
#        "sentence_length_distribution": sentence_length_distribution,
#        "most_common_words": most_common_words,
#        "dicendi_count": dicendi_count,
#        "abstract_words": abstract_count,
#        "concrete_words": concrete_count,
#        "sentence_types": type_counts,
#        "unique_words": len(unique_words),
#        "uniqueness_ratio": uniqueness_ratio,
#        "filler_words": filler_count,
#        "content_words": content_count
#    }

#import pandas as pd
#from collections import defaultdict


#import stanza


##stanza.download("ro")


#nlp = stanza.Pipeline("ro", processors="tokenize,pos,lemma")

#def lemmatize_text_stanza(text):
#    doc = nlp(text)
#    # return " ".join([word.lemma for sent in doc.sentences for word in sent.words])
#    return [word.lemma for sent in doc.sentences for word in sent.words]

#emotion_df = pd.read_csv(
#    "C:/Users/bibis/PycharmProjects/EmotionAnalysis/.venv/data/RoEmoLex_V3_pos (sept2021).csv",
#    delimiter=";"
#)


#emotion_dict = {
#    row["word"]: row[3:-1].astype(float).to_dict()
#    for _, row in emotion_df.iterrows()
#}

#def get_emotion_scores(words):
#    emotions = {col: 0 for col in emotion_df.columns[3:-1]}
#    for word in words:
#        word = word.strip().lower()
#        if word in emotion_dict:
#            for emotion, score in emotion_dict[word].items():
#                emotions[emotion] += score
#    return emotions

#def analyze_emotions(text):
#    sentences = split_into_sentences(text)
#    emotion_over_time = []
#    cumulative_emotions = defaultdict(float)

#    for sentence in sentences:
#        lemmas = lemmatize_text_stanza(sentence)
#        scores = get_emotion_scores(lemmas)
#        emotion_over_time.append({
#            "sentence": sentence,
#            "scores": scores
#        })
#        for emotion, val in scores.items():
#            cumulative_emotions[emotion] += val

#    return {
#        "emotion_over_sentences": emotion_over_time,
#        "total_emotion_distribution": dict(cumulative_emotions)
#    }


#if __name__ == "__main__":
#    if len(sys.argv) < 2:
#        print(json.dumps({"error": "No file provided"}))
#        sys.exit(1)

#    file_path = sys.argv[1]

#    try:
#        text = extract_text_from_docx(file_path)
#        stats = analyze_text(text)
#        emotion_stats = analyze_emotions(text)
#        stats["emotion_analysis"] = emotion_stats
#        print(json.dumps(stats))
#    except Exception as e:
#        print(json.dumps({"error": str(e)}))

import sys
import re
import json
import spacy
from collections import Counter, defaultdict
from docx import Document
import pandas as pd
import stanza

# Load SpaCy for structural analysis
nlp_spacy = spacy.load("ro_core_news_sm")

# Load Stanza for lemmatization
nlp_stanza = stanza.Pipeline("ro", processors="tokenize,pos,lemma", verbose=False)

# Extract text from DOCX
def extract_text_from_docx(file_path):
    doc = Document(file_path)
    return "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])

# Sentence splitting
def split_into_sentences(text):
    return re.split(r'(?<=[.!?]) +', text)

# Load emotion lexicon
emotion_df = pd.read_csv(
    "C:/Users/bibis/PycharmProjects/EmotionAnalysis/.venv/data/RoEmoLex_V3_pos (sept2021).csv",
    delimiter=";"
)
emotion_dict = {
    row["word"]: row[3:-1].astype(float).to_dict()
    for _, row in emotion_df.iterrows()
}

# Get emotion scores for lemmatized words
def get_emotion_scores(words):
    emotions = {col: 0 for col in emotion_df.columns[3:-1]}
    for word in words:
        word = word.strip().lower()
        if word in emotion_dict:
            for emotion, score in emotion_dict[word].items():
                emotions[emotion] += score
    return emotions

# Lemmatize text using Stanza
def lemmatize_text_stanza(text):
    doc = nlp_stanza(text)
    return [word.lemma for sent in doc.sentences for word in sent.words]

# Rolling average for smoothing
def smooth_scores(score_list, window=2):
    df = pd.DataFrame(score_list)
    return df.rolling(window=window, min_periods=1).mean().to_dict(orient="records")

# Analyze text structure (style)
def analyze_text(text):
    doc = nlp_spacy(text)
    sentences = split_into_sentences(text)
    words = re.findall(r'\b\w+\b', text.lower())

    word_count = len(words)
    sentence_count = len(sentences)
    avg_sentence_length = word_count / sentence_count if sentence_count else 0

    sentence_lengths = [len(re.findall(r'\b\w+\b', sentence)) for sentence in sentences]
    sentence_length_distribution = Counter(sentence_lengths).most_common()

    dicendi_verbs = {"conversa", "dialoga", "vorbi", "discuta", "bârfi", "dezbate", "delibera"}
    dicendi_count = sum(1 for token in doc if token.lemma_ in dicendi_verbs and token.pos_ == "VERB")

    abstract_terms = {"iubire", "frică", "libertate", "gândire", "idee", "emoție"}
    abstract_count = 0
    concrete_count = 0
    for token in doc:
        if token.is_alpha and token.pos_ in {"NOUN", "ADJ", "VERB"}:
            if token.lemma_.lower() in abstract_terms:
                abstract_count += 1
            else:
                concrete_count += 1

    type_counts = {"declarative": 0, "interogative": 0, "exclamative": 0}
    for sentence in sentences:
        if sentence.endswith("?"):
            type_counts["interogative"] += 1
        elif sentence.endswith("!"):
            type_counts["exclamative"] += 1
        else:
            type_counts["declarative"] += 1

    filtered_words = [
        token.text.lower() for token in doc
        if token.pos_ not in {"DET", "ADP", "CCONJ", "SCONJ", "PART", "INTJ", "PUNCT", "AUX", "PRON", "ADV"}
    ]
    most_common_words = Counter(filtered_words).most_common(10)

    unique_words = set(words)
    uniqueness_ratio = round(len(unique_words) / word_count, 2) if word_count else 0

    filler_count = sum(1 for token in doc if token.pos_ in {"ADP", "AUX", "SCONJ", "CCONJ"})
    content_count = word_count - filler_count

    return {
        "word_count": word_count,
        "sentence_count": sentence_count,
        "avg_sentence_length": round(avg_sentence_length, 2),
        "sentence_length_distribution": sentence_length_distribution,
        "most_common_words": most_common_words,
        "dicendi_count": dicendi_count,
        "abstract_words": abstract_count,
        "concrete_words": concrete_count,
        "sentence_types": type_counts,
        "unique_words": len(unique_words),
        "uniqueness_ratio": uniqueness_ratio,
        "filler_words": filler_count,
        "content_words": content_count
    }

# Analyze emotions
def analyze_emotions(text):
    sentences = split_into_sentences(text)
    emotion_over_time = []
    cumulative_emotions = defaultdict(float)

    for sentence in sentences:
        lemmas = lemmatize_text_stanza(sentence)
        scores = get_emotion_scores(lemmas)
        emotion_over_time.append({
            "sentence": sentence,
            "scores": scores
        })
        for emotion, val in scores.items():
            cumulative_emotions[emotion] += val

    # Smooth only for visual
    raw_scores = [entry["scores"] for entry in emotion_over_time]
    smoothed_scores = smooth_scores(raw_scores, window=2)

    for idx, entry in enumerate(emotion_over_time):
        entry["smoothed_scores"] = smoothed_scores[idx]

    return {
        "emotion_over_sentences": emotion_over_time,
        "total_emotion_distribution": dict(cumulative_emotions)
    }

# Main entry point
if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(json.dumps({"error": "No file provided"}))
        sys.exit(1)

    file_path = sys.argv[1]

    try:
        text = extract_text_from_docx(file_path)
        stats = analyze_text(text)
        emotion_stats = analyze_emotions(text)
        stats["emotion_analysis"] = emotion_stats
        print(json.dumps(stats))
    except Exception as e:
        print(json.dumps({"error": str(e)}))
